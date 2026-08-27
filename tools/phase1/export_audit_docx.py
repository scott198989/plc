#!/usr/bin/env python3
"""Export the completed Phase 1 adversarial audit Markdown as a reviewable DOCX."""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TEAL = "075F61"
INK = "17212A"
MUTED = "53606C"
LINE = "CBD3D7"
PALE_TEAL = "E8F3F3"
PALE_GRAY = "F3F5F6"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    return parser.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_cell_margins(cell, top: int = 55, start: int = 65, bottom: int = 55, end: int = 65) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, end):
        run._r.append(node)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color, before, after in (
        ("Title", 27, INK, 0, 12),
        ("Subtitle", 12, MUTED, 0, 10),
        ("Heading 1", 19, TEAL, 16, 6),
        ("Heading 2", 14, INK, 13, 5),
        ("Heading 3", 11.5, TEAL, 10, 4),
        ("Heading 4", 9.5, INK, 8, 3),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if "Heading" in name or name == "Title" else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Audit Code" not in styles:
        style = styles.add_style("Audit Code", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Cascadia Mono"
        style.font.size = Pt(7.3)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.14)
        style.paragraph_format.right_indent = Inches(0.14)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(2)

    if "Audit Quote" not in styles:
        style = styles.add_style("Audit Quote", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.italic = True
        style.font.color.rgb = RGBColor.from_string(MUTED)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.right_indent = Inches(0.25)

    header = section.header.paragraphs[0]
    header.text = "PLC Engineering Simulator  |  Phase 1 Adversarial Audit"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(7.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = footer.add_run("Closure candidate — awaiting Scott acceptance  •  ")
    prefix.font.name = "Aptos"
    prefix.font.size = Pt(7.5)
    prefix.font.color.rgb = RGBColor.from_string(MUTED)
    add_page_field(footer)


def split_markdown_row(line: str) -> list[str]:
    text = line.strip()
    if text.startswith("|"):
        text = text[1:]
    if text.endswith("|") and not text.endswith(r"\|"):
        text = text[:-1]
    cells: list[str] = []
    current: list[str] = []
    escaped = False
    for character in text:
        if escaped:
            current.append(character)
            escaped = False
        elif character == "\\":
            escaped = True
        elif character == "|":
            cells.append("".join(current).strip())
            current = []
        else:
            current.append(character)
    if escaped:
        current.append("\\")
    cells.append("".join(current).strip())
    return cells


def is_table_separator(line: str) -> bool:
    cells = split_markdown_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells)


def clean_inline(text: str) -> str:
    value = html.unescape(text)
    value = re.sub(r"<br\s*/?>", "\n", value, flags=re.IGNORECASE)
    value = re.sub(r"<[^>]+>", "", value)
    value = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    value = value.replace("**", "").replace("__", "")
    value = re.sub(r"(?<!\*)\*(?!\*)", "", value)
    value = value.replace("`", "")
    return value.strip()


def add_text_with_breaks(paragraph, text: str, *, bold: bool = False) -> None:
    parts = clean_inline(text).split("\n")
    for index, part in enumerate(parts):
        if index:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(part)
        run.bold = bold


def add_markdown_table(document: Document, raw_rows: list[list[str]]) -> None:
    if not raw_rows:
        return
    column_count = max(len(row) for row in raw_rows)
    table = document.add_table(rows=len(raw_rows), cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, raw_row in enumerate(raw_rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, TEAL)
            elif row_index % 2 == 0:
                set_cell_shading(cell, PALE_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            text = raw_row[column_index] if column_index < len(raw_row) else ""
            add_text_with_breaks(paragraph, text, bold=row_index == 0)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(6.4 if column_count >= 6 else 7.1)
                run.font.color.rgb = RGBColor(255, 255, 255) if row_index == 0 else RGBColor.from_string(INK)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_cover(document: Document, lines: list[str]) -> None:
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(22)
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Phase 1 Adversarial Audit")
    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("Closure candidate — awaiting Scott acceptance")

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(6)
    rule.paragraph_format.space_after = Pt(18)
    run = rule.add_run("▰" * 18)
    run.font.color.rgb = RGBColor.from_string(TEAL)
    run.font.size = Pt(5)

    for line in lines:
        if not line.strip() or line.lstrip().startswith("<!--"):
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        add_text_with_breaks(paragraph, line)
        if paragraph.runs:
            paragraph.runs[0].font.size = Pt(9.5)
    boundary = document.add_paragraph()
    boundary.paragraph_format.space_before = Pt(18)
    boundary.paragraph_format.space_after = Pt(0)
    boundary.add_run(
        "This deliverable reports a Phase 1 closure candidate. It does not constitute Scott's acceptance and does not authorize Phase 2."
    ).bold = True
    document.add_page_break()


def export_markdown(source: Path, output: Path) -> None:
    raw_text = source.read_text(encoding="utf-8")
    if re.search(r"<!--\s*(?:DEFECTS|TASK6|TASK7|NOT_VERIFIABLE|VERDICT)\s*-->", raw_text):
        raise SystemExit("Refusing to export: final-audit placeholder marker remains")

    lines = raw_text.splitlines()
    document = Document()
    configure_document(document)
    properties = document.core_properties
    properties.title = "Phase 1 Adversarial Audit — Closure Candidate"
    properties.subject = "PLC Engineering Simulator Phase 1 closure evidence"
    properties.author = "Codex Automation"
    properties.keywords = "Phase 1, adversarial audit, closure candidate, PLC engineering simulator"

    first_section_index = next(
        (index for index, line in enumerate(lines) if line.startswith("## ")),
        len(lines),
    )
    cover_lines = [line for line in lines[1:first_section_index] if line.strip()]
    add_cover(document, cover_lines)

    index = first_section_index
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_buffer:
            paragraph = document.add_paragraph()
            add_text_with_breaks(paragraph, " ".join(item.strip() for item in paragraph_buffer))
            paragraph_buffer.clear()

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("<!--") and stripped.endswith("-->"):
            flush_paragraph()
            index += 1
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("```"):
            flush_paragraph()
            language = stripped[3:].strip()
            index += 1
            block: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                block.append(lines[index])
                index += 1
            if index < len(lines):
                index += 1
            paragraph = document.add_paragraph(style="Audit Code")
            set_cell = OxmlElement("w:shd")
            set_cell.set(qn("w:fill"), PALE_GRAY)
            paragraph._p.get_or_add_pPr().append(set_cell)
            if language:
                label = paragraph.add_run(f"[{language}]\n")
                label.bold = True
                label.font.color.rgb = RGBColor.from_string(TEAL)
            paragraph.add_run("\n".join(block))
            continue
        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and lines[index + 1].strip().startswith("|")
            and is_table_separator(lines[index + 1])
        ):
            flush_paragraph()
            rows = [split_markdown_row(lines[index])]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_markdown_row(lines[index]))
                index += 1
            add_markdown_table(document, rows)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = min(len(heading.group(1)), 4)
            paragraph = document.add_heading(clean_inline(heading.group(2)), level=level)
            paragraph.paragraph_format.keep_with_next = True
            index += 1
            continue
        if re.match(r"^[-*+]\s+", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_text_with_breaks(paragraph, re.sub(r"^[-*+]\s+", "", stripped))
            index += 1
            continue
        if re.match(r"^\d+[.)]\s+", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            add_text_with_breaks(paragraph, re.sub(r"^\d+[.)]\s+", "", stripped))
            index += 1
            continue
        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = document.add_paragraph(style="Audit Quote")
            add_text_with_breaks(paragraph, stripped.lstrip("> "))
            index += 1
            continue
        if re.fullmatch(r"[-*_]{3,}", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph()
            run = paragraph.add_run("—" * 48)
            run.font.color.rgb = RGBColor.from_string(LINE)
            index += 1
            continue

        paragraph_buffer.append(line)
        index += 1

    flush_paragraph()
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"DOCX_PATH={output.resolve()}")
    print(f"DOCX_BYTES={output.stat().st_size}")


def main() -> int:
    args = parse_args()
    source = args.input.resolve(strict=True)
    output = args.output.resolve(strict=False)
    export_markdown(source, output)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
